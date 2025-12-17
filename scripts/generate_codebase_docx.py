#!/usr/bin/env python3
"""
Generate a DOCX file from the morph-magic-web codebase with table of contents
and FBLA rubric highlights.

Requires: pip install python-docx

Open in Word and right-click the TOC > Update Field to populate page numbers.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent

INCLUDE_PATTERNS = [
    "src/pages/*.tsx",
    "src/components/Navbar.tsx",
    "src/components/Layout.tsx",
    "src/components/PageTransition.tsx",
    "src/components/GlassCard.tsx",
    "src/components/ReCaptcha.tsx",
    "src/components/FloatingOrbs.tsx",
    "src/components/MorphScene.tsx",
    "src/components/MorphBlob.tsx",
    "src/components/TiltCard.tsx",
    "src/components/MagneticButton.tsx",
    "src/components/ScrollAnimations.tsx",
    "src/components/ChatWidget.tsx",
    "src/components/GuidedTour.tsx",
    "src/hooks/useAuth.ts",
    "src/hooks/useFavorites.ts",
    "src/hooks/useUserReviews.ts",
    "src/hooks/AuthContext.tsx",
    "src/hooks/use-toast.ts",
    "vite.config.ts",
    "tailwind.config.ts",
]

# FBLA rubric criteria mapping: file path pattern -> callout text (highlighted before code)
RUBRIC_HIGHLIGHTS = {
    "src/hooks/useFavorites.ts": (
        "FBLA Rubric: Prompt — Saving/bookmarking favorites; Data storage (localStorage per user); "
        "Program structure."
    ),
    "src/hooks/useUserReviews.ts": (
        "FBLA Rubric: Prompt — Users leave reviews/ratings; Data storage (localStorage); "
        "Dynamic updating of reviews."
    ),
    "src/hooks/useAuth.ts": (
        "FBLA Rubric: Data storage; Program structure; Usability & results."
    ),
    "src/hooks/AuthContext.tsx": (
        "FBLA Rubric: Program structure; Conciseness; Modularity."
    ),
    "src/pages/SubmitBusiness.tsx": (
        "FBLA Rubric: Input validation (Zod schema); Bot verification (CAPTCHA); "
        "Commentary quality; Form handling."
    ),
    "src/pages/BusinessLogin.tsx": (
        "FBLA Rubric: Bot verification (CAPTCHA); Usability; Navigation."
    ),
    "src/components/ReCaptcha.tsx": (
        "FBLA Rubric: Prompt — Verification to prevent bots; Program structure."
    ),
    "src/pages/Directory.tsx": (
        "FBLA Rubric: Prompt — Sort by category, reviews, ratings; Usability & navigation; "
        "Logical program sequence; Dynamic filtering."
    ),
    "src/pages/BusinessDetail.tsx": (
        "FBLA Rubric: Prompt — Users leave reviews; Sort by reviews/ratings display; "
        "Usability; Output reports."
    ),
    "src/pages/Index.tsx": (
        "FBLA Rubric: Prompt — Display deals/coupons; Usability; Navigation; "
        "Output reports quality."
    ),
    "src/pages/MyFavorites.tsx": (
        "FBLA Rubric: Prompt — Saving/bookmarking favorites; Usability; "
        "Results accuracy and dynamic updating."
    ),
    "src/components/Navbar.tsx": (
        "FBLA Rubric: Usability & navigation; Program usability features."
    ),
    "src/components/Layout.tsx": (
        "FBLA Rubric: Logical program sequence; Conciseness; Navigation structure."
    ),
}


def gather_files() -> list[Path]:
    files = []
    for pattern in INCLUDE_PATTERNS:
        p = Path(ROOT) / pattern
        if "*" in pattern:
            parent = p.parent
            glob_pat = p.name
            files.extend(sorted(parent.glob(glob_pat)))
        else:
            if p.exists() and p.is_file():
                files.append(p)
    exclude = {"node_modules", "package.json", "package-lock.json"}
    return [f for f in files if "node_modules" not in str(f) and f.name not in exclude]


def add_toc_field(doc: Document):
    """Insert a Word TOC field. User must Update Field in Word."""
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    for tag, val in [
        ("w:fldChar", "begin"),
        ("w:instrText", 'TOC \\o "1-3" \\h \\z \\u'),
        ("w:fldChar", "separate"),
        ("w:fldChar", "end"),
    ]:
        if tag == "w:instrText":
            e = OxmlElement("w:instrText")
            e.set(qn("xml:space"), "preserve")
            e.text = val
            r.append(e)
        else:
            e = OxmlElement("w:fldChar")
            e.set(qn("w:fldCharType"), val)
            r.append(e)


def add_rubric_callout(doc: Document, text: str):
    """Add a highlighted paragraph for FBLA rubric criteria. Uses GRAY_50 for B&W printing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.highlight_color = WD_COLOR_INDEX.GRAY_50


def add_heading_with_indicator(doc: Document, text: str, level: int, has_rubric: bool):
    """Add heading; prefix with ◆ (diamond) if rubric-relevant. Prints clearly in B&W."""
    heading_text = "◆  " + text if has_rubric else text
    doc.add_heading(heading_text, level=level)


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    doc.add_heading("Cypress LocalLink — Codebase Documentation", 0)
    doc.add_paragraph(
        "This document contains the main source code for the Cypress LocalLink "
        "local business directory web application, with FBLA rubric highlights."
    )
    doc.add_paragraph()

    # FBLA Rubric Reference
    doc.add_heading("FBLA Rubric Reference (2025–2026 Coding & Programming)", level=1)
    doc.add_paragraph("Program Readability (60 pts):", style="List Bullet")
    doc.add_paragraph("Appropriate identifiers; Commentary quality; Documentation completeness.", style="List Bullet 2")
    doc.add_paragraph("Program Structure & Content (100 pts):", style="List Bullet")
    doc.add_paragraph("Conciseness; Data storage with security; Logical sequence; Usability & navigation.", style="List Bullet 2")
    doc.add_paragraph("Usability & Results (40 pts):", style="List Bullet")
    doc.add_paragraph("Results accuracy and dynamic updating; Output reports quality.", style="List Bullet 2")
    doc.add_paragraph()
    doc.add_paragraph(
        "Byte-Sized Business Boost prompt: Sort by category; Users leave reviews/ratings; "
        "Sort by reviews/ratings; Favorites; Deals/coupons; Bot verification."
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("◆ = Rubric-relevant code (gray shading marks criteria)")
    r.font.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph()

    # Example SQL Tables
    doc.add_heading("Example SQL Tables", level=1)
    doc.add_paragraph(
        "The app uses static data and localStorage. Below are example SQL table definitions "
        "that could support a backend implementation."
    )
    doc.add_paragraph()

    sql_example = """-- Example: businesses table (could support the directory)
CREATE TABLE businesses (
  id          VARCHAR(50) PRIMARY KEY,
  name        VARCHAR(100) NOT NULL,
  category    VARCHAR(50) NOT NULL,
  rating      DECIMAL(2,1) DEFAULT 0,
  review_count INT DEFAULT 0,
  address     VARCHAR(200) NOT NULL,
  description TEXT,
  image       VARCHAR(500),
  price_range VARCHAR(10),
  lat         DECIMAL(10,7),
  lng         DECIMAL(10,7),
  created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Example: user_reviews table (matches useUserReviews localStorage)
CREATE TABLE user_reviews (
  id          VARCHAR(50) PRIMARY KEY,
  business_id VARCHAR(50) REFERENCES businesses(id),
  user_id     VARCHAR(50) NOT NULL,
  author      VARCHAR(100) NOT NULL,
  rating      INT CHECK (rating >= 1 AND rating <= 5),
  text        TEXT NOT NULL,
  created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

-- Example: favorites table (matches useFavorites localStorage)
CREATE TABLE favorites (
  user_id     VARCHAR(50) NOT NULL,
  business_id VARCHAR(50) REFERENCES businesses(id),
  PRIMARY KEY (user_id, business_id)
);"""
    add_code_block(doc, sql_example)
    doc.add_page_break()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "Right-click the gray area below and choose Update Field (or Ctrl+A, F9) to populate page numbers."
    )
    add_toc_field(doc)
    doc.add_page_break()

    # Source Code with rubric highlights
    doc.add_heading("Source Code", level=1)
    files = sorted(gather_files(), key=lambda x: str(x.relative_to(ROOT)))

    for f in files:
        rel = f.relative_to(ROOT)
        rel_str = rel.as_posix()
        has_rubric = rel_str in RUBRIC_HIGHLIGHTS

        # Add heading with ◆ indicator for rubric-relevant files
        add_heading_with_indicator(doc, rel_str, level=2, has_rubric=has_rubric)

        # Add rubric callout if this file maps to rubric criteria
        if has_rubric:
            add_rubric_callout(doc, RUBRIC_HIGHLIGHTS[rel_str])

        # Add code
        try:
            content = f.read_text(encoding="utf-8", errors="replace")
            add_code_block(doc, content)
        except Exception as e:
            doc.add_paragraph(f"[Error reading file: {e}]")
        doc.add_paragraph()

    out = ROOT / "Cypress-LocalLink-Codebase.docx"
    doc.save(out)
    print(f"Saved: {out}")
    print(f"Included {len(files)} files with {len(RUBRIC_HIGHLIGHTS)} rubric highlights.")


if __name__ == "__main__":
    main()
